"""
FastAPI application for the EagleEye security intelligence platform.
Provides REST API endpoints for all system capabilities.
"""

import io
import json
import logging
import math
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Optional

import numpy as np
import pandas as pd
from dateutil import parser as dateparser
from docx import Document as DocxDocument
from fastapi import FastAPI, Form, HTTPException, Query, UploadFile, File
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

logger = logging.getLogger(__name__)

from eagleeye.alerts.alert_engine import AlertEngine
from eagleeye.alerts.report_generator import ReportGenerator
from eagleeye.analysis.correlation import CorrelationAnalyzer
from eagleeye.analysis.network import NetworkAnalyzer
from eagleeye.analysis.spatial import SpatialAnalyzer
from eagleeye.analysis.temporal import TemporalAnalyzer
from eagleeye.core.config import TARGET_STATES
from eagleeye.core.database import init_db
from eagleeye.data.processor import DataProcessor
from eagleeye.data.schemas import IncidentCreate, IntelReportCreate
from eagleeye.intel.news_monitor import NewsMonitor
from eagleeye.llm.claude_engine import engine as claude_engine
from eagleeye.models.anomaly_detector import AnomalyDetector
from eagleeye.models.predictor import BanditActivityPredictor
from eagleeye.models.threat_scorer import ThreatScorer

def _sanitize(obj):
    """Convert numpy/pandas types to native Python types for JSON serialization."""
    if isinstance(obj, dict):
        return {k: _sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize(v) for v in obj]
    if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
        return None
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        v = float(obj)
        return None if math.isnan(v) or math.isinf(v) else v
    if isinstance(obj, np.ndarray):
        return _sanitize(obj.tolist())
    if isinstance(obj, np.bool_):
        return bool(obj)
    return obj


app = FastAPI(
    title="EagleEye",
    description="AI-powered security intelligence platform for predicting and "
                "anticipating armed bandit activities in Northern Nigeria.",
    version="0.1.0",
)

# Mount static files for the web UI
_web_dir = Path(__file__).resolve().parent.parent / "web"
app.mount("/static", StaticFiles(directory=str(_web_dir / "static")), name="static")

# Register web page routes
from eagleeye.web.routes import router as web_router  # noqa: E402
app.include_router(web_router)

# Global state
predictor = BanditActivityPredictor()
anomaly_detector = AnomalyDetector()


@app.on_event("startup")
def startup():
    init_db()


# --- Incident Management ---

@app.post("/api/incidents", tags=["Incidents"])
def create_incident(incident: IncidentCreate):
    """Record a new security incident."""
    processor = DataProcessor()
    try:
        stored = processor.ingest_incident(incident)
        return {"id": stored.id, "status": "created", "state": stored.state}
    finally:
        processor.close()


@app.get("/api/incidents", tags=["Incidents"])
def list_incidents(
    state: Optional[str] = None,
    incident_type: Optional[str] = None,
    days: int = Query(default=30, ge=1, le=365),
):
    """List incidents with optional filters."""
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(
            state=state, incident_type=incident_type, start_date=start_date
        )
        if df.empty:
            return {"incidents": [], "total": 0}
        records = _sanitize(df.to_dict(orient="records"))
        return {
            "incidents": records,
            "total": len(df),
        }
    finally:
        processor.close()


# --- Intelligence Reports ---

@app.post("/api/intel/report", tags=["Intelligence"])
def submit_intel_report(report: IntelReportCreate):
    """Submit a raw intelligence report for parsing and processing."""
    processor = DataProcessor()
    monitor = NewsMonitor()
    try:
        # Save raw report text for future reprocessing
        saved_report = processor.ingest_intel_report(
            title=report.title or "",
            content=report.content,
            source=report.source or "",
            source_url=report.source_url or "",
            published_date=report.published_date,
            region=report.region,
        )

        incidents = monitor.process_single_report(
            text=report.content,
            source=report.source or "",
            source_url=report.source_url or "",
            published_date=report.published_date,
        )

        # Update report with extraction count
        saved_report.processed = 1
        saved_report.extracted_incidents_count = len(incidents)
        processor.session.commit()

        return {
            "status": "processed",
            "report_id": saved_report.id,
            "incidents_extracted": len(incidents),
            "incidents": incidents,
        }
    finally:
        processor.close()
        monitor.close()


@app.post("/api/intel/upload", tags=["Intelligence"])
async def upload_intel_document(file: UploadFile = File(...)):
    """Upload a Word document (.docx) or text file for NLP parsing and incident extraction."""
    filename = file.filename or "uploaded_file"
    suffix = Path(filename).suffix.lower()

    if suffix not in (".docx", ".txt"):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{suffix}'. Upload .docx or .txt files.",
        )

    raw_bytes = await file.read()

    # Extract text from the document
    if suffix == ".docx":
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to parse .docx file '{filename}': {e}")
            raise HTTPException(status_code=400, detail=f"Could not read .docx file: {e}")
    else:
        text = raw_bytes.decode("utf-8", errors="replace")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Document is empty — no text extracted.")

    # Try to extract a date from the document header/footer
    published_date = None
    try:
        snippet = (text[:500] + " " + text[-500:]) if len(text) > 1000 else text
        import re
        date_match = re.search(
            r'(\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s*,?\s*\d{4})',
            snippet, re.IGNORECASE)
        if not date_match:
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', snippet)
        if not date_match:
            date_match = re.search(r'(\d{1,2}/\d{1,2}/\d{4})', snippet)
        if date_match:
            published_date = dateparser.parse(date_match.group(1), fuzzy=True, dayfirst=True)
    except Exception:
        pass

    # Save raw text and process through the NLP pipeline
    processor = DataProcessor()
    monitor = NewsMonitor()
    try:
        # Save raw report for future reprocessing
        saved_report = processor.ingest_intel_report(
            title=filename,
            content=text,
            source=filename,
            published_date=published_date,
        )

        incidents = monitor.process_single_report(
            text=text,
            source=filename,
            source_url="",
            published_date=published_date,
        )

        # Update report with extraction count
        saved_report.processed = 1
        saved_report.extracted_incidents_count = len(incidents)
        processor.session.commit()

        return {
            "status": "processed",
            "report_id": saved_report.id,
            "filename": filename,
            "text_length": len(text),
            "incidents_extracted": len(incidents),
            "incidents": incidents,
        }
    finally:
        processor.close()
        monitor.close()


@app.post("/api/intel/collect", tags=["Intelligence"])
def run_news_collection():
    """Trigger a news collection cycle from configured sources."""
    monitor = NewsMonitor()
    try:
        result = monitor.run_collection_cycle()
        return result
    except Exception as e:
        logger.error(f"News collection failed: {e}")
        return {"status": "error", "message": str(e), "stats": monitor.stats}
    finally:
        monitor.close()


@app.get("/api/intel/reports", tags=["Intelligence"])
def list_intel_reports():
    """List all stored intel reports."""
    from eagleeye.core.database import get_session, IntelReport
    session = get_session()
    try:
        reports = session.query(IntelReport).order_by(IntelReport.ingested_at.desc()).all()
        return {
            "reports": [
                {
                    "id": r.id,
                    "title": r.title,
                    "source": r.source,
                    "content_length": len(r.content) if r.content else 0,
                    "processed": r.processed,
                    "extracted_incidents_count": r.extracted_incidents_count,
                    "ingested_at": r.ingested_at.isoformat() if r.ingested_at else None,
                }
                for r in reports
            ],
            "total": len(reports),
        }
    finally:
        session.close()


@app.post("/api/intel/reprocess", tags=["Intelligence"])
def reprocess_all_reports(clear_old: bool = True):
    """Reprocess all stored intel reports through the AI parser (Ollama/Claude/HF).

    Args:
        clear_old: If True (default), delete all existing incidents before reprocessing.
    """
    from eagleeye.core.database import get_session, IntelReport, Incident

    # Phase 1: Read all report data and clear old incidents, then close session
    # This avoids holding a read lock while the monitor writes new incidents.
    session = get_session()
    try:
        reports_raw = session.query(IntelReport).all()
        if not reports_raw:
            return {"status": "no_reports", "message": "No stored reports to reprocess."}

        # Snapshot report data into plain dicts
        report_data = []
        for r in reports_raw:
            report_data.append({
                "id": r.id,
                "title": r.title,
                "content": r.content,
                "source": r.source or "",
                "source_url": r.source_url or "",
                "published_date": r.published_date,
            })

        deleted_count = 0
        if clear_old:
            deleted_count = session.query(Incident).delete()
            session.commit()
            logger.info(f"Cleared {deleted_count} old incidents before reprocessing")
    finally:
        session.close()

    # Phase 2: Reprocess each report (monitor uses its own session for writes)
    monitor = NewsMonitor()
    total_new_incidents = 0
    results = []

    try:
        for i, rd in enumerate(report_data, 1):
            logger.info(f"Reprocessing report {i}/{len(report_data)}: {rd['title'] or rd['source'] or 'untitled'}")
            incidents = monitor.process_single_report(
                text=rd["content"],
                source=rd["source"],
                source_url=rd["source_url"],
                published_date=rd["published_date"],
            )
            total_new_incidents += len(incidents)
            results.append({
                "report_id": rd["id"],
                "title": rd["title"],
                "incidents_extracted": len(incidents),
            })
    finally:
        monitor.close()

    # Phase 3: Update report statuses in a separate session
    session2 = get_session()
    try:
        for res in results:
            report = session2.query(IntelReport).get(res["report_id"])
            if report:
                report.processed = 1
                report.extracted_incidents_count = res["incidents_extracted"]
        session2.commit()
    finally:
        session2.close()

    return {
        "status": "completed",
        "reports_reprocessed": len(report_data),
        "old_incidents_cleared": deleted_count,
        "total_new_incidents": total_new_incidents,
        "active_ai_backend": claude_engine.active_backend,
        "results": results,
    }


# --- Analysis ---

@app.get("/api/analysis/temporal", tags=["Analysis"])
def temporal_analysis(state: Optional[str] = None, days: int = 180):
    """Run temporal pattern analysis."""
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(state=state, start_date=start_date)
        if df.empty:
            return {"status": "no_data"}

        analyzer = TemporalAnalyzer(df)
        result = {}
        try:
            result["frequency"] = analyzer.frequency_analysis(state, window_days=days)
        except Exception as e:
            logger.error(f"Temporal frequency_analysis failed: {e}")
            result["frequency"] = {"status": "error", "message": str(e)}
        try:
            result["seasonal_pattern"] = analyzer.seasonal_pattern(state)
        except Exception as e:
            logger.error(f"Temporal seasonal_pattern failed: {e}")
            result["seasonal_pattern"] = {"status": "error"}
        try:
            result["day_of_week"] = analyzer.day_of_week_pattern(state)
        except Exception as e:
            logger.error(f"Temporal day_of_week_pattern failed: {e}")
            result["day_of_week"] = {"status": "error"}
        try:
            result["surges"] = analyzer.detect_surge_periods(state=state)
        except Exception as e:
            logger.error(f"Temporal detect_surge_periods failed: {e}")
            result["surges"] = []
        try:
            result["next_window_prediction"] = analyzer.predict_next_window(state)
        except Exception as e:
            logger.error(f"Temporal predict_next_window failed: {e}")
            result["next_window_prediction"] = {"status": "error"}

        result = _sanitize(result)

        # Claude narration
        try:
            narrative = claude_engine.narrate_analysis("temporal", result)
            if narrative:
                result["narrative"] = narrative
        except Exception as e:
            logger.debug(f"Claude temporal narration failed: {e}")

        return result
    finally:
        processor.close()


@app.get("/api/analysis/spatial", tags=["Analysis"])
def spatial_analysis(state: Optional[str] = None, days: int = 180):
    """Run geographic/spatial pattern analysis."""
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(state=state, start_date=start_date)
        if df.empty:
            return {"status": "no_data"}

        analyzer = SpatialAnalyzer(df)
        result = {}
        try:
            result["hotspots"] = analyzer.hotspot_analysis(state)
        except Exception as e:
            logger.error(f"Spatial hotspot_analysis failed: {e}")
            result["hotspots"] = []
        try:
            result["state_comparison"] = analyzer.state_comparison()
        except Exception as e:
            logger.error(f"Spatial state_comparison failed: {e}")
            result["state_comparison"] = []
        try:
            result["clusters"] = analyzer.geographic_clustering()
        except Exception as e:
            logger.error(f"Spatial geographic_clustering failed: {e}")
            result["clusters"] = []
        try:
            result["cross_border"] = analyzer.cross_border_analysis()
        except Exception as e:
            logger.error(f"Spatial cross_border_analysis failed: {e}")
            result["cross_border"] = {}

        result = _sanitize(result)

        try:
            narrative = claude_engine.narrate_analysis("spatial", result)
            if narrative:
                result["narrative"] = narrative
        except Exception as e:
            logger.debug(f"Claude spatial narration failed: {e}")

        return result
    finally:
        processor.close()


@app.get("/api/analysis/correlation", tags=["Analysis"])
def correlation_analysis(state: Optional[str] = None, days: int = 365):
    """Run correlation and pattern analysis."""
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(state=state, start_date=start_date)
        if df.empty:
            return {"status": "no_data"}

        analyzer = CorrelationAnalyzer(df)
        result = {}
        try:
            result["attack_sequences"] = analyzer.attack_sequence_patterns(state=state)
        except Exception as e:
            logger.error(f"Correlation attack_sequence_patterns failed: {e}")
            result["attack_sequences"] = []
        try:
            result["escalation_indicators"] = analyzer.escalation_indicators(state=state)
        except Exception as e:
            logger.error(f"Correlation escalation_indicators failed: {e}")
            result["escalation_indicators"] = {}
        try:
            result["tactic_correlations"] = analyzer.tactic_correlation_matrix(state=state)
        except Exception as e:
            logger.error(f"Correlation tactic_correlation_matrix failed: {e}")
            result["tactic_correlations"] = {}
        try:
            result["retaliatory_patterns"] = analyzer.retaliatory_pattern_detection(state=state)
        except Exception as e:
            logger.error(f"Correlation retaliatory_pattern_detection failed: {e}")
            result["retaliatory_patterns"] = {}
        try:
            result["vulnerability_windows"] = analyzer.vulnerability_windows(state=state)
        except Exception as e:
            logger.error(f"Correlation vulnerability_windows failed: {e}")
            result["vulnerability_windows"] = {}

        result = _sanitize(result)

        try:
            narrative = claude_engine.narrate_analysis("correlation", result)
            if narrative:
                result["narrative"] = narrative
        except Exception as e:
            logger.debug(f"Claude correlation narration failed: {e}")

        return result
    finally:
        processor.close()


@app.get("/api/analysis/network", tags=["Analysis"])
def network_analysis(state: Optional[str] = None, days: int = 365):
    """Run network and group analysis."""
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(state=state, start_date=start_date)
        if df.empty:
            return {"status": "no_data"}

        analyzer = NetworkAnalyzer(df)
        result = {}
        try:
            result["territories"] = analyzer.territory_mapping()
        except Exception as e:
            logger.error(f"Network territory_mapping failed: {e}")
            result["territories"] = []
        try:
            result["operational_profiles"] = analyzer.operational_pattern_profiling()
        except Exception as e:
            logger.error(f"Network operational_pattern_profiling failed: {e}")
            result["operational_profiles"] = []
        try:
            result["supply_routes"] = analyzer.supply_route_analysis()
        except Exception as e:
            logger.error(f"Network supply_route_analysis failed: {e}")
            result["supply_routes"] = {}
        try:
            result["activity_timeline"] = analyzer.group_activity_timeline(state)
        except Exception as e:
            logger.error(f"Network group_activity_timeline failed: {e}")
            result["activity_timeline"] = {}
        try:
            result["operational_lulls"] = analyzer.identify_operational_lulls(state)
        except Exception as e:
            logger.error(f"Network identify_operational_lulls failed: {e}")
            result["operational_lulls"] = []

        result = _sanitize(result)

        try:
            narrative = claude_engine.narrate_analysis("network", result)
            if narrative:
                result["narrative"] = narrative
        except Exception as e:
            logger.debug(f"Claude network narration failed: {e}")

        return result
    finally:
        processor.close()


# --- Threat Scoring ---

@app.get("/api/threat/scores", tags=["Threat Assessment"])
def get_threat_scores():
    """Get current threat scores for all target states."""
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            return {"status": "no_data", "scores": []}
        scorer = ThreatScorer(df)
        return {"scores": _sanitize(scorer.score_all_states())}
    finally:
        processor.close()


@app.get("/api/threat/state/{state}", tags=["Threat Assessment"])
def get_state_threat(state: str):
    """Get detailed threat assessment for a specific state."""
    if state not in TARGET_STATES:
        raise HTTPException(status_code=400, detail=f"Invalid state. Must be one of: {TARGET_STATES}")
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            return {"status": "no_data"}
        scorer = ThreatScorer(df)
        return _sanitize(scorer.score_state(state))
    finally:
        processor.close()


# --- Predictions ---

@app.post("/api/predict/train", tags=["Predictions"])
def train_model():
    """Train the prediction model on all available historical data."""
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            raise HTTPException(status_code=400, detail="No incident data available for training")
        result = predictor.train(df)

        # Auto-generate alerts from analysis
        alerts_generated = 0
        alert_engine = AlertEngine()
        try:
            for state in TARGET_STATES:
                try:
                    state_df = df[df["state"].str.lower() == state.lower()]
                    if state_df.empty:
                        continue

                    # Surge detection
                    temporal = TemporalAnalyzer(state_df)
                    surges = temporal.detect_surge_periods(state=state)
                    if isinstance(surges, list):
                        for surge in surges[-2:]:  # Latest 2 surges only
                            alert_engine.generate_surge_alert(surge, state=state)
                            alerts_generated += 1

                    # Escalation detection
                    correlation = CorrelationAnalyzer(state_df)
                    escalation = correlation.escalation_indicators(state=state)
                    if isinstance(escalation, dict) and escalation.get("escalation_score", 0) >= 4:
                        alert_engine.generate_escalation_alert(escalation, state)
                        alerts_generated += 1
                except Exception as e:
                    logger.warning(f"Alert generation for {state} failed: {e}")
        finally:
            alert_engine.close()

        result["alerts_generated"] = alerts_generated
        return result
    finally:
        processor.close()


@app.post("/api/predict/forecast", tags=["Predictions"])
def generate_forecast(
    state: Optional[str] = None,
    horizon_days: int = Query(default=14, ge=1, le=90),
):
    """Generate activity predictions for the next N days."""
    if not predictor.is_trained:
        raise HTTPException(status_code=400, detail="Model not trained. Call /api/predict/train first.")

    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe(state=state)

        # Build current context from recent data
        now = datetime.now()
        recent_7d = df[df["date"] >= now - timedelta(days=7)] if not df.empty else df
        recent_30d = df[df["date"] >= now - timedelta(days=30)] if not df.empty else df

        # Type-specific momentum: count of each type in last 30 days
        type_momentum = {}
        if not recent_30d.empty and "incident_type" in recent_30d.columns:
            type_counts = recent_30d["incident_type"].value_counts()
            for t in type_counts.index:
                type_momentum[t] = int(type_counts[t])

        context = {
            "date": now,
            "state": state or "Zamfara",
            "recent_incidents_7d": len(recent_7d),
            "recent_incidents_14d": len(df[df["date"] >= now - timedelta(days=14)]) if not df.empty else 0,
            "recent_incidents_30d": len(recent_30d),
            "recent_killed_7d": int(recent_7d["casualties_killed"].sum()) if not recent_7d.empty else 0,
            "recent_killed_30d": int(recent_30d["casualties_killed"].sum()) if not recent_30d.empty else 0,
            "recent_incidents_in_state_30d": len(recent_30d),
            "type_momentum_30d": type_momentum,
        }

        predictions = predictor.predict_multi_day(context, horizon_days)

        # Auto-generate alerts for high-threat predictions
        alerts_generated = 0
        alert_engine = AlertEngine()
        try:
            for pred in predictions:
                if pred.get("threat_level", 0) >= 4:
                    alert_engine.generate_prediction_alert(pred)
                    alerts_generated += 1
                    if alerts_generated >= 5:
                        break
        except Exception as e:
            logger.warning(f"Prediction alert generation failed: {e}")
        finally:
            alert_engine.close()

        return {
            "state": state or "all",
            "horizon_days": horizon_days,
            "predictions": predictions,
            "alerts_generated": alerts_generated,
        }
    finally:
        processor.close()


# --- Anomaly Detection ---

@app.post("/api/anomaly/fit", tags=["Anomaly Detection"])
def fit_anomaly_detector():
    """Fit the anomaly detector on historical data."""
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            raise HTTPException(status_code=400, detail="No data available")
        result = anomaly_detector.fit(df)
        return result
    finally:
        processor.close()


@app.get("/api/anomaly/detect", tags=["Anomaly Detection"])
def detect_anomalies(days: int = Query(default=30, ge=1, le=365)):
    """Detect anomalies in recent incident data."""
    if not anomaly_detector.is_fitted:
        raise HTTPException(status_code=400, detail="Anomaly detector not fitted. Call /api/anomaly/fit first.")
    processor = DataProcessor()
    try:
        start_date = datetime.now() - timedelta(days=days)
        df = processor.get_incidents_dataframe(start_date=start_date)
        if df.empty:
            return {"anomalies": []}
        anomalies = anomaly_detector.detect(df)

        # Auto-generate alerts for anomalies (capped at 10)
        alerts_generated = 0
        if anomalies:
            alert_engine = AlertEngine()
            try:
                for anomaly in anomalies[:10]:
                    try:
                        alert_engine.generate_anomaly_alert(anomaly)
                        alerts_generated += 1
                    except Exception as e:
                        logger.warning(f"Anomaly alert generation failed: {e}")
            finally:
                alert_engine.close()

        return {"anomalies": anomalies, "total": len(anomalies),
                "alerts_generated": alerts_generated}
    finally:
        processor.close()


# --- Alerts ---

@app.get("/api/alerts", tags=["Alerts"])
def get_alerts(state: Optional[str] = None, min_severity: int = 1):
    """Get active security alerts."""
    engine = AlertEngine()
    try:
        alerts = engine.get_active_alerts(state=state, min_severity=min_severity)
        return {"alerts": alerts, "total": len(alerts)}
    finally:
        engine.close()


@app.post("/api/alerts/{alert_id}/acknowledge", tags=["Alerts"])
def acknowledge_alert(alert_id: int, acknowledged_by: str = "operator"):
    """Acknowledge a security alert."""
    engine = AlertEngine()
    try:
        success = engine.acknowledge_alert(alert_id, acknowledged_by)
        if not success:
            raise HTTPException(status_code=404, detail="Alert not found")
        return {"status": "acknowledged"}
    finally:
        engine.close()


# --- Reports ---

@app.get("/api/reports/sitrep", tags=["Reports"])
def generate_sitrep(state: Optional[str] = None, days: int = 30):
    """Generate a Situation Report (SITREP)."""
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            return {"status": "no_data"}
        generator = ReportGenerator(df)
        return _sanitize(generator.generate_situation_report(period_days=days, state=state))
    finally:
        processor.close()


@app.get("/api/reports/threat-briefing/{state}", tags=["Reports"])
def generate_threat_briefing(state: str):
    """Generate a focused threat briefing for a state."""
    if state not in TARGET_STATES:
        raise HTTPException(status_code=400, detail=f"Invalid state. Must be one of: {TARGET_STATES}")
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            return {"status": "no_data"}
        generator = ReportGenerator(df)
        return _sanitize(generator.generate_threat_briefing(state))
    finally:
        processor.close()


class TemplatedSitrepRequest(BaseModel):
    template: str
    state: Optional[str] = None
    days: int = 30
    context_texts: Optional[List[str]] = None


class SaveTemplateRequest(BaseModel):
    template: str


class AreaThreatBriefingRequest(BaseModel):
    area_name: Optional[str] = None
    state: str
    lat: Optional[float] = None
    lon: Optional[float] = None


class AiEnginesRequest(BaseModel):
    engines: List[str]


@app.post("/api/reports/sitrep/templated", tags=["Reports"])
def generate_templated_sitrep(req: TemplatedSitrepRequest):
    """Generate a SITREP following a user-provided template format."""
    if not req.template or not req.template.strip():
        raise HTTPException(status_code=400, detail="Template text is required")
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            return {"status": "no_data"}
        generator = ReportGenerator(df)
        result = generator.generate_templated_sitrep(
            template=req.template, period_days=req.days, state=req.state,
            context_texts=req.context_texts,
        )
        return _sanitize(result)
    finally:
        processor.close()


@app.post("/api/reports/threat-briefing/area", tags=["Reports"])
def generate_area_threat_briefing(req: AreaThreatBriefingRequest):
    """Generate a terrain-integrated area threat briefing using JIFC framework."""
    if not req.area_name and (req.lat is None or req.lon is None):
        raise HTTPException(
            status_code=400,
            detail="Provide either area_name or both lat and lon coordinates",
        )
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        if df.empty:
            df = pd.DataFrame(columns=[
                "id", "date", "state", "lga", "incident_type",
                "casualties_killed", "casualties_injured", "kidnapped_count",
            ])
        generator = ReportGenerator(df)
        result = generator.generate_area_threat_briefing(
            area_name=req.area_name or "",
            state=req.state,
            lat=req.lat,
            lon=req.lon,
        )
        return _sanitize(result)
    finally:
        processor.close()


@app.get("/api/reports/predictive-outlook", tags=["Reports"])
def generate_predictive_outlook(
    state: Optional[str] = None,
    horizon_days: int = 14,
):
    """Generate a predictive outlook report."""
    if not predictor.is_trained:
        raise HTTPException(status_code=400, detail="Model not trained")

    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe(state=state)
        now = datetime.now()
        recent_7d = df[df["date"] >= now - timedelta(days=7)] if not df.empty else df

        context = {
            "date": now,
            "state": state or "Zamfara",
            "recent_incidents_7d": len(recent_7d),
            "recent_incidents_30d": len(df),
            "recent_killed_7d": int(recent_7d["casualties_killed"].sum()) if not recent_7d.empty else 0,
        }

        predictions = predictor.predict_multi_day(context, horizon_days)
        generator = ReportGenerator(df)
        return _sanitize(generator.generate_predictive_outlook(predictions, state))
    finally:
        processor.close()


# --- Settings ---

@app.get("/api/settings/sitrep-template", tags=["Settings"])
def get_sitrep_template():
    """Get the saved SITREP template."""
    from eagleeye.core.database import get_session, AppSetting
    session = get_session()
    try:
        setting = session.query(AppSetting).filter_by(key="sitrep_template").first()
        if setting:
            return {
                "template": setting.value,
                "updated_at": setting.updated_at.isoformat() if setting.updated_at else None,
            }
        return {"template": None}
    finally:
        session.close()


@app.put("/api/settings/sitrep-template", tags=["Settings"])
def save_sitrep_template(req: SaveTemplateRequest):
    """Save or update the default SITREP template."""
    if not req.template or not req.template.strip():
        raise HTTPException(status_code=400, detail="Template text is required")
    from eagleeye.core.database import get_session, AppSetting
    session = get_session()
    try:
        setting = session.query(AppSetting).filter_by(key="sitrep_template").first()
        if setting:
            setting.value = req.template
            setting.updated_at = datetime.utcnow()
        else:
            setting = AppSetting(key="sitrep_template", value=req.template)
            session.add(setting)
        session.commit()
        return {"status": "saved"}
    finally:
        session.close()


@app.delete("/api/settings/sitrep-template", tags=["Settings"])
def delete_sitrep_template():
    """Delete the saved SITREP template."""
    from eagleeye.core.database import get_session, AppSetting
    session = get_session()
    try:
        setting = session.query(AppSetting).filter_by(key="sitrep_template").first()
        if setting:
            session.delete(setting)
            session.commit()
        return {"status": "deleted"}
    finally:
        session.close()


@app.get("/api/settings/ai-engines", tags=["Settings"])
def get_ai_engines():
    """Get the current AI engine priority configuration."""
    from eagleeye.core.database import get_session, AppSetting
    from eagleeye.llm.claude_engine import VALID_ENGINES, DEFAULT_ENGINE_ORDER
    session = get_session()
    try:
        setting = session.query(AppSetting).filter_by(key="ai_engines").first()
        if setting:
            engines = json.loads(setting.value)
            engines = [e for e in engines if e in VALID_ENGINES]
            if engines:
                return {"engines": engines}
        return {"engines": list(DEFAULT_ENGINE_ORDER)}
    finally:
        session.close()


@app.put("/api/settings/ai-engines", tags=["Settings"])
def save_ai_engines(req: AiEnginesRequest):
    """Save the AI engine priority configuration."""
    from eagleeye.core.database import get_session, AppSetting
    from eagleeye.llm.claude_engine import VALID_ENGINES

    # Validate
    if not req.engines:
        raise HTTPException(status_code=400, detail="At least one engine must be selected")
    for e in req.engines:
        if e not in VALID_ENGINES:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid engine '{e}'. Must be one of: {', '.join(VALID_ENGINES)}",
            )

    session = get_session()
    try:
        setting = session.query(AppSetting).filter_by(key="ai_engines").first()
        value = json.dumps(req.engines)
        if setting:
            setting.value = value
            setting.updated_at = datetime.utcnow()
        else:
            setting = AppSetting(key="ai_engines", value=value)
            session.add(setting)
        session.commit()
    finally:
        session.close()

    # Apply immediately without restart
    claude_engine.reload_config()
    return {"status": "saved", "engines": req.engines}


@app.post("/api/reports/sitrep/upload-context", tags=["Reports"])
async def upload_sitrep_context(files: List[UploadFile] = File(...)):
    """Upload context reports (.docx, .txt) for SITREP generation."""
    results = []
    for file in files:
        filename = file.filename or "uploaded_file"
        suffix = Path(filename).suffix.lower()

        if suffix not in (".docx", ".txt"):
            results.append({
                "filename": filename,
                "error": f"Unsupported file type '{suffix}'. Use .docx or .txt.",
            })
            continue

        raw_bytes = await file.read()

        if suffix == ".docx":
            try:
                doc = DocxDocument(io.BytesIO(raw_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text = "\n\n".join(paragraphs)
            except Exception as e:
                results.append({"filename": filename, "error": f"Could not read .docx: {e}"})
                continue
        else:
            text = raw_bytes.decode("utf-8", errors="replace")

        if not text.strip():
            results.append({"filename": filename, "error": "Document is empty"})
            continue

        results.append({
            "filename": filename,
            "text": text,
            "length": len(text),
        })

    return {"files": results}


@app.post("/api/reports/analyze", tags=["Reports"])
async def analyze_uploaded_report(file: UploadFile = File(...)):
    """Upload a report (.docx, .txt, .pdf) and get comprehensive AI analysis."""
    filename = file.filename or "uploaded_file"
    suffix = Path(filename).suffix.lower()

    if suffix not in (".docx", ".txt", ".pdf"):
        return {"error": f"Unsupported file type '{suffix}'. Use .docx, .txt, or .pdf."}

    raw_bytes = await file.read()

    if suffix == ".docx":
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            return {"error": f"Could not read .docx: {e}"}
    elif suffix == ".pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            text = "\n\n".join(pages)
        except ImportError:
            return {"error": "PDF support requires pdfplumber. Install it with: pip install pdfplumber"}
        except Exception as e:
            return {"error": f"Could not read .pdf: {e}"}
    else:
        text = raw_bytes.decode("utf-8", errors="replace")

    if not text.strip():
        return {"error": "Document is empty"}

    analysis = claude_engine.analyze_report(text)
    if not analysis:
        return {"error": "AI analysis unavailable. Check that an AI engine is configured in Settings."}

    return {
        "filename": filename,
        "text_length": len(text),
        "analysis": analysis,
        "ai_backend": claude_engine.active_backend,
    }


@app.post("/api/reports/sitrep/from-report", tags=["Reports"])
async def generate_sitrep_from_report(
    file: UploadFile = File(...),
    template: Optional[str] = Form(None),
):
    """Generate a SITREP based exclusively on an uploaded report."""
    filename = file.filename or "uploaded_file"
    suffix = Path(filename).suffix.lower()

    if suffix not in (".docx", ".txt", ".pdf"):
        return {"error": f"Unsupported file type '{suffix}'. Use .docx, .txt, or .pdf."}

    raw_bytes = await file.read()

    if suffix == ".docx":
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
        except Exception as e:
            return {"error": f"Could not read .docx: {e}"}
    elif suffix == ".pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
            text = "\n\n".join(pages)
        except ImportError:
            return {"error": "PDF support requires pdfplumber. Install with: pip install pdfplumber"}
        except Exception as e:
            return {"error": f"Could not read .pdf: {e}"}
    else:
        text = raw_bytes.decode("utf-8", errors="replace")

    if not text.strip():
        return {"error": "Document is empty"}

    tmpl = template.strip() if template else None
    narrative = claude_engine.generate_sitrep_from_report(text, template=tmpl)
    if not narrative:
        return {"error": "AI engine unavailable. Check that an AI engine is configured in Settings."}

    return {
        "filename": filename,
        "text_length": len(text),
        "narrative": narrative,
        "ai_backend": claude_engine.active_backend,
    }


# --- System ---

@app.get("/api/intelligence-brief", tags=["Intelligence"])
def get_intelligence_brief():
    """Generate an AI-powered daily intelligence brief for the dashboard."""
    processor = DataProcessor()
    try:
        df = processor.get_incidents_dataframe()
        now = datetime.now()

        # Build context
        context = {
            "date": now.strftime("%A, %d %B %Y"),
            "total_incidents": len(df),
        }

        if not df.empty:
            df_copy = df.copy()
            df_copy["date"] = pd.to_datetime(df_copy["date"])
            recent_7d = df_copy[df_copy["date"] >= now - timedelta(days=7)]
            recent_30d = df_copy[df_copy["date"] >= now - timedelta(days=30)]

            context.update({
                "incidents_last_7d": len(recent_7d),
                "incidents_last_30d": len(recent_30d),
                "killed_last_30d": int(recent_30d["casualties_killed"].sum()),
                "kidnapped_last_30d": int(recent_30d["kidnapped_count"].sum()),
                "states_affected_30d": recent_30d["state"].nunique() if not recent_30d.empty else 0,
                "top_states_30d": recent_30d.groupby("state")["id"].count().nlargest(3).to_dict() if not recent_30d.empty else {},
                "top_incident_types_30d": recent_30d.groupby("incident_type")["id"].count().nlargest(5).to_dict() if not recent_30d.empty else {},
            })

            # Threat scores
            try:
                scorer = ThreatScorer(df_copy)
                scores = scorer.score_all_states()
                context["threat_scores"] = [
                    {"state": s["state"], "level": s.get("threat_label", "LOW"),
                     "score": s.get("overall_score", 0)}
                    for s in scores[:3]
                ]
            except Exception:
                pass

            context["model_trained"] = predictor.is_trained

        # Try Claude first
        brief = claude_engine.generate_intelligence_brief(context)

        if not brief:
            # Basic fallback template
            inc_7 = context.get("incidents_last_7d", 0)
            inc_30 = context.get("incidents_last_30d", 0)
            killed = context.get("killed_last_30d", 0)
            brief = (
                f"EagleEye Intelligence Brief — {context['date']}\n\n"
                f"Over the past 30 days, {inc_30} security incidents have been recorded "
                f"across the Northwest region, with {inc_7} incidents in the last 7 days. "
                f"Fatalities stand at {killed} for the period.\n\n"
                f"{'The prediction model is operational and generating forecasts.' if context.get('model_trained') else 'The prediction model has not yet been trained.'}\n\n"
                f"Full analysis available on the Analysis and Reports pages."
            )

        return {
            "brief": brief,
            "generated_at": now.isoformat(),
            "ai_generated": claude_engine.any_ai_available,
            "ai_backend": claude_engine.active_backend,
        }
    finally:
        processor.close()


@app.get("/api/status", tags=["System"])
def system_status():
    """Get system status and configuration."""
    return {
        "status": "operational",
        "version": "0.1.0",
        "target_states": TARGET_STATES,
        "model_trained": predictor.is_trained,
        "anomaly_detector_fitted": anomaly_detector.is_fitted,
        "model_stats": predictor.training_stats if predictor.is_trained else None,
        "ollama_available": claude_engine.is_available,
        "any_ai_available": claude_engine.any_ai_available,
        "active_ai_backend": claude_engine.active_backend,
        "engines_config": claude_engine._load_enabled_engines(),
        "engines_status": claude_engine.get_engine_status(),
    }
